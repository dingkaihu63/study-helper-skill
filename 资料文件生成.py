"""
专业 Word 文档生成器
支持：封面、目录、多级标题、中英文字体区分、表格、图片、页眉页脚
"""

import os
from dataclasses import dataclass, field
from typing import List, Optional, Tuple, Dict, Any, Union
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ============ 配置类 ============

@dataclass
class DocumentConfig:
    """文档全局配置"""
    # 页面设置
    page_width: Cm = field(default_factory=lambda: Cm(21.0))  # A4 宽度
    page_height: Cm = field(default_factory=lambda: Cm(29.7))  # A4 高度
    margin_top: Cm = field(default_factory=lambda: Cm(2.54))
    margin_bottom: Cm = field(default_factory=lambda: Cm(2.54))
    margin_left: Cm = field(default_factory=lambda: Cm(3.17))  # 装订边稍宽
    margin_right: Cm = field(default_factory=lambda: Cm(3.17))

    # 中文字体
    chinese_font: str = "宋体"
    chinese_font_heading: str = "黑体"

    # 英文字体
    english_font: str = "Times New Roman"
    english_font_heading: str = "Times New Roman"

    # 字号设置 (磅值)
    font_size_main_title: Pt = field(default_factory=lambda: Pt(22))  # 主标题：二号
    font_size_title1: Pt = field(default_factory=lambda: Pt(16))  # 一级标题：三号
    font_size_title2: Pt = field(default_factory=lambda: Pt(14))  # 二级标题：四号
    font_size_title3: Pt = field(default_factory=lambda: Pt(12))  # 三级标题：小四加粗
    font_size_body: Pt = field(default_factory=lambda: Pt(12))  # 正文：小四
    font_size_table: Pt = field(default_factory=lambda: Pt(10.5))  # 表格：五号
    font_size_caption: Pt = field(default_factory=lambda: Pt(10.5))  # 题注：五号

    # 行距
    line_spacing_body: float = 1.5  # 正文1.5倍行距
    line_spacing_heading: float = 1.5  # 标题行距

    # 颜色
    color_main_title: RGBColor = field(default_factory=lambda: RGBColor(0x00, 0x00, 0x00))
    color_heading1: RGBColor = field(default_factory=lambda: RGBColor(0x00, 0x33, 0x66))  # 深蓝
    color_heading2: RGBColor = field(default_factory=lambda: RGBColor(0x00, 0x66, 0x99))  # 中蓝
    color_heading3: RGBColor = field(default_factory=lambda: RGBColor(0x33, 0x33, 0x33))  # 深灰
    color_body: RGBColor = field(default_factory=lambda: RGBColor(0x33, 0x33, 0x33))

    # 段落间距
    space_before_heading1: Pt = field(default_factory=lambda: Pt(24))
    space_after_heading1: Pt = field(default_factory=lambda: Pt(12))
    space_before_heading2: Pt = field(default_factory=lambda: Pt(18))
    space_after_heading2: Pt = field(default_factory=lambda: Pt(6))
    space_before_body: Pt = field(default_factory=lambda: Pt(0))
    space_after_body: Pt = field(default_factory=lambda: Pt(6))

    # 封面设置
    cover_title: str = ""
    cover_subtitle: str = ""
    cover_author: str = ""
    cover_date: str = field(default_factory=lambda: datetime.now().strftime("%Y-%m-%d"))
    cover_organization: str = ""


# ============ 字体工具 ============

class FontManager:
    """字体管理器，处理中英文字体混合"""

    def __init__(self, config: DocumentConfig):
        self.config = config

    def set_run_font(
            self,
            run,
            font_size: Pt = None,
            bold: bool = False,
            color: RGBColor = None,
            is_chinese: bool = True
    ):
        """设置 run 的字体属性"""
        font = run.font

        # 字号
        if font_size:
            font.size = font_size

        # 粗细
        font.bold = bold

        # 颜色
        if color:
            font.color.rgb = color

        # 字体设置（关键：分别设置中西文字体）
        if is_chinese:
            font.name = self.config.chinese_font
            run._element.rPr.rFonts.set(qn('w:eastAsia'), self.config.chinese_font)
        else:
            font.name = self.config.english_font
            run._element.rPr.rFonts.set(qn('w:eastAsia'), self.config.chinese_font)

    def apply_mixed_font(
            self,
            paragraph,
            text: str,
            font_size: Pt = None,
            bold: bool = False,
            color: RGBColor = None
    ):
        """
        智能识别中英文并分别设置字体
        中文用宋体，英文/数字用 Times New Roman
        """
        import re

        # 正则分割：中文字符 vs 非中文字符
        pattern = re.compile(
            r'([\u4e00-\u9fff\u3000-\u303f\uff00-\uffef]+|[^\u4e00-\u9fff\u3000-\u303f\uff00-\uffef]+)')
        segments = pattern.findall(text)

        for segment in segments:
            run = paragraph.add_run(segment)
            is_chinese = bool(re.match(r'[\u4e00-\u9fff\u3000-\u303f\uff00-\uffef]', segment))

            self.set_run_font(
                run,
                font_size=font_size,
                bold=bold,
                color=color,
                is_chinese=is_chinese
            )

        return paragraph


# ============ 核心生成器 ============

class WordDocumentBuilder:
    """Word 文档构建器"""

    def __init__(self, config: Optional[DocumentConfig] = None):
        self.config = config or DocumentConfig()
        self.doc = Document()
        self.font_mgr = FontManager(self.config)
        self._heading_counters = {1: 0, 2: 0, 3: 0}
        self._setup_document()
        self._setup_styles()

    def _setup_document(self):
        """设置页面布局"""
        section = self.doc.sections[0]
        section.page_width = self.config.page_width
        section.page_height = self.config.page_height
        section.top_margin = self.config.margin_top
        section.bottom_margin = self.config.margin_bottom
        section.left_margin = self.config.margin_left
        section.right_margin = self.config.margin_right

    def _setup_styles(self):
        """预设文档样式"""
        # 正文样式
        style = self.doc.styles['Normal']
        style.font.name = self.config.english_font
        style.font.size = self.config.font_size_body
        style._element.rPr.rFonts.set(qn('w:eastAsia'), self.config.chinese_font)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        style.paragraph_format.line_spacing = self.config.line_spacing_body
        style.paragraph_format.space_after = self.config.space_after_body
        style.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进2字符

    def _set_paragraph_format(
            self,
            paragraph,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            space_before=None,
            space_after=None,
            line_spacing=None,
            first_line_indent=None
    ):
        """统一设置段落格式"""
        pf = paragraph.paragraph_format
        pf.alignment = alignment

        if space_before is not None:
            pf.space_before = space_before
        if space_after is not None:
            pf.space_after = space_after
        if line_spacing is not None:
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            pf.line_spacing = line_spacing
        if first_line_indent is not None:
            pf.first_line_indent = first_line_indent

    # ============ 封面 ============

    def add_cover(self):
        """添加专业封面"""
        config = self.config

        # 顶部留白
        for _ in range(3):
            p = self.doc.add_paragraph()
            self._set_paragraph_format(p, space_after=Pt(12))

        # 机构名称
        if config.cover_organization:
            p = self.doc.add_paragraph()
            self._set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(24))
            self.font_mgr.apply_mixed_font(
                p, config.cover_organization,
                font_size=Pt(16), bold=True, color=RGBColor(0x00, 0x33, 0x66)
            )

        # 主标题
        p = self.doc.add_paragraph()
        self._set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(60), space_after=Pt(12))
        self.font_mgr.apply_mixed_font(
            p, config.cover_title or "文档标题",
            font_size=config.font_size_main_title, bold=True, color=config.color_main_title
        )

        # 副标题
        if config.cover_subtitle:
            p = self.doc.add_paragraph()
            self._set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(48))
            self.font_mgr.apply_mixed_font(
                p, config.cover_subtitle,
                font_size=Pt(14), bold=False, color=RGBColor(0x66, 0x66, 0x66)
            )

        # 分隔线效果（用下划线段落模拟）
        p = self.doc.add_paragraph()
        self._set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(36))
        run = p.add_run("━" * 30)
        run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
        run.font.size = Pt(8)

        # 作者信息
        info_items = [
            ("编制单位", config.cover_organization),
            ("编 制 人", config.cover_author),
            ("编制日期", config.cover_date),
        ]

        for label, value in info_items:
            if value:
                p = self.doc.add_paragraph()
                self._set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(8))
                self.font_mgr.apply_mixed_font(
                    p, f"{label}：{value}",
                    font_size=Pt(12), bold=False, color=RGBColor(0x33, 0x33, 0x33)
                )

        # 分页
        self.doc.add_page_break()

    # ============ 目录 ============

    def add_toc(self, title: str = "目  录"):
        """添加自动目录"""
        # 目录标题
        p = self.doc.add_paragraph()
        self._set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(24))
        self.font_mgr.apply_mixed_font(
            p, title,
            font_size=self.config.font_size_title1, bold=True, color=self.config.color_heading1
        )

        # 使用 Word 域代码插入自动目录
        # 注意：需要手动更新域（在 Word 中右键更新域）
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run()
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._r.append(fldChar1)

        run2 = paragraph.add_run()
        instrText = parse_xml(
            f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
        run2._r.append(instrText)

        run3 = paragraph.add_run()
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
        run3._r.append(fldChar2)

        # 占位提示
        run4 = paragraph.add_run("[请右键更新域以生成目录]")
        run4.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run4.font.size = Pt(10)

        run5 = paragraph.add_run()
        fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run5._r.append(fldChar3)

        self.doc.add_page_break()

    # ============ 标题 ============

    def add_heading1(self, text: str):
        """一级标题（章标题）"""
        self._heading_counters[1] += 1
        self._heading_counters[2] = 0
        self._heading_counters[3] = 0

        num = f"第{self._number_to_chinese(self._heading_counters[1])}章 "
        full_text = num + text

        p = self.doc.add_paragraph()
        self._set_paragraph_format(
            p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
            space_before=self.config.space_before_heading1,
            space_after=self.config.space_after_heading1,
            line_spacing=self.config.line_spacing_heading,
            first_line_indent=None
        )

        self.font_mgr.apply_mixed_font(
            p, full_text,
            font_size=self.config.font_size_title1,
            bold=True,
            color=self.config.color_heading1
        )

        # 添加段落下边框（装饰线）
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="0066CC"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)

    def add_heading2(self, text: str):
        """二级标题"""
        self._heading_counters[2] += 1
        self._heading_counters[3] = 0

        num = f"{self._heading_counters[1]}.{self._heading_counters[2]} "
        full_text = num + text

        p = self.doc.add_paragraph()
        self._set_paragraph_format(
            p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
            space_before=self.config.space_before_heading2,
            space_after=self.config.space_after_heading2,
            line_spacing=self.config.line_spacing_heading,
            first_line_indent=None
        )

        self.font_mgr.apply_mixed_font(
            p, full_text,
            font_size=self.config.font_size_title2,
            bold=True,
            color=self.config.color_heading2
        )

    def add_heading3(self, text: str):
        """三级标题"""
        self._heading_counters[3] += 1

        num = f"{self._heading_counters[1]}.{self._heading_counters[2]}.{self._heading_counters[3]} "
        full_text = num + text

        p = self.doc.add_paragraph()
        self._set_paragraph_format(
            p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
            space_before=Pt(12),
            space_after=Pt(6),
            line_spacing=self.config.line_spacing_heading,
            first_line_indent=None
        )

        self.font_mgr.apply_mixed_font(
            p, full_text,
            font_size=self.config.font_size_title3,
            bold=True,
            color=self.config.color_heading3
        )

    def _number_to_chinese(self, num: int) -> str:
        """数字转中文（简化版）"""
        chinese_nums = ['零', '一', '二', '三', '四', '五', '六', '七', '八', '九', '十']
        if num <= 10:
            return chinese_nums[num]
        elif num < 20:
            return '十' + (chinese_nums[num - 10] if num > 10 else '')
        else:
            return str(num)

    # ============ 正文 ============

    def add_paragraph(self, text: str, bold: bool = False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
        """
        添加正文段落
        自动识别中英文并分别设置字体
        """
        p = self.doc.add_paragraph()
        self._set_paragraph_format(
            p, alignment=alignment,
            space_before=self.config.space_before_body,
            space_after=self.config.space_after_body,
            line_spacing=self.config.line_spacing_body,
            first_line_indent=Cm(0.74)  # 首行缩进
        )

        self.font_mgr.apply_mixed_font(
            p, text,
            font_size=self.config.font_size_body,
            bold=bold,
            color=self.config.color_body
        )
        return p

    def add_bullet_list(self, items: List[str]):
        """添加项目符号列表"""
        for item in items:
            p = self.doc.add_paragraph(style='List Bullet')
            self._set_paragraph_format(
                p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                space_after=Pt(4),
                line_spacing=self.config.line_spacing_body,
                first_line_indent=None
            )
            self.font_mgr.apply_mixed_font(
                p, item,
                font_size=self.config.font_size_body,
                color=self.config.color_body
            )

    def add_numbered_list(self, items: List[str]):
        """添加编号列表"""
        for item in items:
            p = self.doc.add_paragraph(style='List Number')
            self._set_paragraph_format(
                p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                space_after=Pt(4),
                line_spacing=self.config.line_spacing_body,
                first_line_indent=None
            )
            self.font_mgr.apply_mixed_font(
                p, item,
                font_size=self.config.font_size_body,
                color=self.config.color_body
            )

    # ============ 表格 ============

    def add_table(
            self,
            headers: List[str],
            rows: List[List[str]],
            col_widths: Optional[List[Cm]] = None,
            style_name: str = "Table Grid"
    ):
        """
        添加美观的表格

        Args:
            headers: 表头
            rows: 数据行
            col_widths: 列宽列表（可选）
            style_name: 表格样式
        """
        table = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = style_name
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 设置表格自动调整
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')

        # 设置表格宽度为页面宽度的 100%
        width = self.config.page_width - self.config.margin_left - self.config.margin_right
        tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
        tblPr.append(tblW)

        # 设置列宽
        if col_widths:
            for i, width in enumerate(col_widths):
                for cell in table.columns[i].cells:
                    tcPr = cell._tc.get_or_add_tcPr()
                    tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{int(width.cm * 567)}" w:type="dxa"/>')
                    tcPr.append(tcW)

        # 填充表头
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            cell = header_cells[i]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            # 清除默认段落
            cell.paragraphs[0].clear()
            p = cell.paragraphs[0]
            self._set_paragraph_format(
                p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                space_before=Pt(4), space_after=Pt(4),
                first_line_indent=None
            )

            self.font_mgr.apply_mixed_font(
                p, header,
                font_size=self.config.font_size_table,
                bold=True,
                color=RGBColor(0xFF, 0xFF, 0xFF)
            )

            # 表头背景色（深蓝）
            shading = parse_xml(
                f'<w:shd {nsdecls("w")} w:fill="0066CC" w:val="clear"/>'
            )
            cell._tc.get_or_add_tcPr().append(shading)

        # 填充数据行
        for row_idx, row_data in enumerate(rows):
            row_cells = table.rows[row_idx + 1].cells
            for col_idx, cell_text in enumerate(row_data):
                cell = row_cells[col_idx]
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

                cell.paragraphs[0].clear()
                p = cell.paragraphs[0]
                self._set_paragraph_format(
                    p, alignment=WD_ALIGN_PARAGRAPH.CENTER if col_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT,
                    space_before=Pt(3), space_after=Pt(3),
                    first_line_indent=None
                )

                # 斑马纹效果
                if row_idx % 2 == 1:
                    shading = parse_xml(
                        f'<w:shd {nsdecls("w")} w:fill="F2F7FC" w:val="clear"/>'
                    )
                    cell._tc.get_or_add_tcPr().append(shading)

                self.font_mgr.apply_mixed_font(
                    p, str(cell_text),
                    font_size=self.config.font_size_table,
                    color=self.config.color_body
                )

        # 表格后留白
        self.doc.add_paragraph()
        return table

    # ============ 图片与题注 ============

    def add_picture(self, image_path: str, width: Optional[Cm] = None, caption: str = ""):
        """添加图片及题注"""
        if not os.path.exists(image_path):
            p = self.doc.add_paragraph()
            self.font_mgr.apply_mixed_font(p, f"[图片未找到: {image_path}]")
            return

        # 计算合适宽度
        if width is None:
            available_width = self.config.page_width - self.config.margin_left - self.config.margin_right
            width = available_width * 0.8

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_path, width=width)

        # 题注
        if caption:
            cap_p = self.doc.add_paragraph()
            self._set_paragraph_format(
                cap_p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                space_before=Pt(4), space_after=Pt(12),
                first_line_indent=None
            )
            self.font_mgr.apply_mixed_font(
                cap_p, f"图 {caption}",
                font_size=self.config.font_size_caption,
                color=RGBColor(0x66, 0x66, 0x66)
            )

    # ============ 分页与分节 ============

    def add_page_break(self):
        """分页"""
        self.doc.add_page_break()

    def add_section_break(self):
        """分节（用于不同页眉页脚）"""
        self.doc.add_section()

    # ============ 页眉页脚 ============

    def set_header(self, text: str, alignment=WD_ALIGN_PARAGRAPH.CENTER):
        """设置页眉"""
        section = self.doc.sections[0]
        header = section.header
        header.is_linked_to_previous = False

        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.clear()
        p.alignment = alignment

        # 添加下划线
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="CCCCCC"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)

        self.font_mgr.apply_mixed_font(
            p, text,
            font_size=Pt(9), color=RGBColor(0x99, 0x99, 0x99)
        )

    def set_footer(self, text: str = None, alignment=WD_ALIGN_PARAGRAPH.CENTER):
        """设置页脚（含页码）"""
        section = self.doc.sections[0]
        footer = section.footer
        footer.is_linked_to_previous = False

        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.clear()
        p.alignment = alignment

        if text:
            self.font_mgr.apply_mixed_font(
                p, text,
                font_size=Pt(9), color=RGBColor(0x99, 0x99, 0x99)
            )
        else:
            # 插入页码域
            run1 = p.add_run("第 ")
            run1.font.size = Pt(9)
            run1.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

            # PAGE 域
            run2 = p.add_run()
            fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
            run2._r.append(fldChar1)

            run3 = p.add_run()
            instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
            run3._r.append(instrText)

            run4 = p.add_run()
            fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
            run4._r.append(fldChar2)

            run5 = p.add_run("1")
            run5.font.size = Pt(9)

            run6 = p.add_run()
            fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
            run6._r.append(fldChar3)

            run7 = p.add_run(" 页 / 共 ")
            run7.font.size = Pt(9)
            run7.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

            # NUMPAGES 域
            run8 = p.add_run()
            fldChar4 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
            run8._r.append(fldChar4)

            run9 = p.add_run()
            instrText2 = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> NUMPAGES </w:instrText>')
            run9._r.append(instrText2)

            run10 = p.add_run()
            fldChar5 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
            run10._r.append(fldChar5)

            run11 = p.add_run("1")
            run11.font.size = Pt(9)

            run12 = p.add_run()
            fldChar6 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
            run12._r.append(fldChar6)

            run13 = p.add_run(" 页")
            run13.font.size = Pt(9)
            run13.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # ============ 保存 ============

    def save(self, output_path: str):
        """保存文档"""
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_path))
        print(f"文档已保存: {output_path}")
        return str(output_path)


# ============ 便捷使用类 ============

class QuickDocument:
    """快速创建标准文档的便捷类"""

    def __init__(self, title: str, author: str = "", organization: str = ""):
        config = DocumentConfig(
            cover_title=title,
            cover_author=author,
            cover_organization=organization,
            cover_date=datetime.now().strftime("%Y-%m-%d")
        )
        self.builder = WordDocumentBuilder(config)
        self._has_cover = False
        self._has_toc = False

    def generate_full_document(
            self,
            sections: List[Dict[str, Any]],
            output_path: str,
            include_cover: bool = True,
            include_toc: bool = True,
            header_text: str = None,
            footer_page_number: bool = True
    ):
        """
        生成完整文档

        Args:
            sections: 内容区块列表，格式如下：
                [
                    {
                        "level": 1,           # 标题级别 1/2/3
                        "title": "章节标题",
                        "content": [          # 内容列表
                            {"type": "text", "text": "正文内容..."},
                            {"type": "table", "headers": [...], "rows": [...]},
                            {"type": "list", "items": [...], "ordered": False},
                        ]
                    }
                ]
        """
        # 封面
        if include_cover and not self._has_cover:
            self.builder.add_cover()
            self._has_cover = True

        # 目录
        if include_toc and not self._has_toc:
            self.builder.add_toc()
            self._has_toc = True

        # 页眉页脚
        if header_text:
            self.builder.set_header(header_text)
        if footer_page_number:
            self.builder.set_footer()

        # 内容
        for section in sections:
            level = section.get("level", 1)
            title = section.get("title", "")

            # 添加标题
            if level == 1:
                self.builder.add_heading1(title)
            elif level == 2:
                self.builder.add_heading2(title)
            elif level == 3:
                self.builder.add_heading3(title)

            # 添加内容
            for content_item in section.get("content", []):
                ctype = content_item.get("type", "text")

                if ctype == "text":
                    self.builder.add_paragraph(content_item.get("text", ""))

                elif ctype == "table":
                    self.builder.add_table(
                        headers=content_item.get("headers", []),
                        rows=content_item.get("rows", []),
                        col_widths=content_item.get("col_widths")
                    )

                elif ctype == "list":
                    items = content_item.get("items", [])
                    if content_item.get("ordered", False):
                        self.builder.add_numbered_list(items)
                    else:
                        self.builder.add_bullet_list(items)

                elif ctype == "page_break":
                    self.builder.add_page_break()

        # 保存
        return self.builder.save(output_path)


# ============ 演示 ============

def demo():
    """生成示例文档"""

    doc = QuickDocument(
        title="基于深度学习的机械臂故障检测系统研究报告",
        author="张三",
        organization="智能装备研究院"
    )

    sections = [
        {
            "level": 1,
            "title": "绪论",
            "content": [
                {"type": "text",
                 "text": "随着工业4.0和智能制造的快速发展，工业机械臂作为自动化生产线的核心装备，其运行可靠性直接关系到生产效率和产品质量。传统的故障检测方法主要依赖人工巡检和定期维护，存在检测效率低、漏检率高、无法实时预警等问题。"},
                {"type": "text",
                 "text": "近年来，深度学习技术在图像识别、自然语言处理等领域取得了突破性进展，为机械臂故障检测提供了新的技术路径。本研究旨在探索基于LSTM（Long Short-Term Memory）网络的机械臂故障检测方法，实现对机械臂运行状态的实时监测与智能诊断。"},
            ]
        },
        {
            "level": 1,
            "title": "相关技术综述",
            "content": [
                {"type": "text",
                 "text": "本章将介绍机械臂故障检测领域的相关技术，包括传统信号处理方法、机器学习方法和深度学习方法，并分析各类方法的优缺点。"},
                {"type": "page_break"},
            ]
        },
        {
            "level": 2,
            "title": "机械臂故障类型分析",
            "content": [
                {"type": "text",
                 "text": "机械臂在长期运行过程中，由于磨损、疲劳、腐蚀等因素，会出现多种类型的故障。根据故障发生的部位和机理，可将机械臂故障分为以下几类："},
                {"type": "list", "items": [
                    "关节故障：包括电机故障、减速器磨损、编码器失效等",
                    "连杆故障：包括结构变形、裂纹、连接松动等",
                    "末端执行器故障：包括夹持力不足、定位精度下降等",
                    "控制系统故障：包括传感器漂移、通信中断、软件异常等"
                ], "ordered": False},
                {"type": "text", "text": "表2-1列出了各类故障的典型特征参数及其检测难度评估。"},
                {
                    "type": "table",
                    "headers": ["故障类型", "特征参数", "检测方法", "难度等级"],
                    "rows": [
                        ["关节电机故障", "电流、温度、振动", "电流分析法", "中等"],
                        ["减速器磨损", "振动、噪声、背隙", "振动频谱分析", "较难"],
                        ["编码器失效", "位置误差、脉冲丢失", "冗余校验", "较易"],
                        ["连杆裂纹", "应变、模态频率", "模态分析", "困难"],
                        ["传感器漂移", "数据偏差、异常值", "统计检验", "中等"],
                    ],
                    "col_widths": [Cm(3.5), Cm(3.5), Cm(3.5), Cm(2.5)]
                },
            ]
        },
        {
            "level": 2,
            "title": "深度学习模型对比",
            "content": [
                {"type": "text",
                 "text": "为了验证不同深度学习模型在机械臂故障检测任务上的性能，本文设计了三种模型架构进行对比实验：基线LSTM模型、LSTM+Attention模型和LSTM+Attention+Uncertainty模型。"},
                {
                    "type": "table",
                    "headers": ["模型架构", "参数量", "准确率(%)", "推理速度(ms)", "F1-Score"],
                    "rows": [
                        ["Baseline LSTM", "1.2M", "87.3", "12.5", "0.854"],
                        ["LSTM + Attention", "1.8M", "91.6", "18.2", "0.903"],
                        ["LSTM + Attn + Uncertainty", "2.1M", "93.8", "22.4", "0.927"],
                    ],
                    "col_widths": [Cm(4), Cm(2.5), Cm(2.5), Cm(2.5), Cm(2)]
                },
                {"type": "text",
                 "text": "实验结果表明，引入注意力机制和不确定性量化的模型在各项指标上均优于基线模型，特别是在处理复杂故障模式时表现出更强的鲁棒性。"},
            ]
        },
        {
            "level": 1,
            "title": "实验验证与分析",
            "content": [
                {"type": "text",
                 "text": "为验证所提方法的有效性，本章在UR3协作机械臂平台上进行了大量实验。实验数据采集自真实工业场景，涵盖了正常状态和多种故障状态。"},
                {"type": "text",
                 "text": "实验环境配置如下：机械臂型号为Universal Robots UR3，控制柜版本为CB3.1，采样频率为500Hz。数据采集包括六维力/力矩信号、关节电流、关节位置和电机温度等关键参数。"},
                {"type": "text",
                 "text": "图3-1展示了实验平台的整体架构。系统主要由数据采集模块、信号预处理模块、特征提取模块和故障诊断模块组成。"},
            ]
        },
        {
            "level": 1,
            "title": "结论与展望",
            "content": [
                {"type": "text",
                 "text": "本文针对工业机械臂故障检测问题，提出了一种基于LSTM网络和注意力机制的智能诊断方法，并引入了不确定性量化机制以提高模型的可靠性。"},
                {"type": "text", "text": "主要研究成果包括："},
                {"type": "list", "items": [
                    "构建了包含多种故障类型的机械臂数据集，为后续研究提供了数据基础",
                    "设计了融合注意力机制的LSTM网络架构，有效提升了故障特征提取能力",
                    "引入不确定性量化方法，增强了模型在边缘样本上的判断可靠性",
                    "在UR3机械臂上完成了实验验证，证明了方法的有效性"
                ], "ordered": True},
                {"type": "text",
                 "text": "未来工作将聚焦于多模态数据融合、在线学习机制以及边缘计算部署等方面，以进一步提升系统的实用性和泛化能力。"},
            ]
        },
    ]

    output_path = "机械臂故障检测报告.docx"
    doc.generate_full_document(
        sections=sections,
        output_path=output_path,
        include_cover=True,
        include_toc=True,
        header_text="基于深度学习的机械臂故障检测系统研究报告",
        footer_page_number=True
    )

    print(f"\n✅ 文档生成完成: {os.path.abspath(output_path)}")
    print("\n📋 使用说明：")
    print("1. 打开文档后，右键目录区域选择「更新域」→「更新整个目录」")
    print("2. 页码显示为「1」时，选中页脚按 F9 更新域")
    print("3. 如需调整样式，可在 Word 中修改对应样式集")


if __name__ == "__main__":
    demo()